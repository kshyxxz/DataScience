from docx import Document
import re
import os

file_path = "unicode_input.docx"

if not os.path.exists(file_path):
    print("File not found:", file_path)
else:
    doc = Document(file_path)
    print("File loaded successfully!")


def reorder_text(text):
    # Move 'ि' before consonant
    text = re.sub(r'(.)(ि)', r'\2\1', text)
    return text


def unicode_to_preeti(text):
    text = reorder_text(text)

    # Handle combined chars FIRST
    text = text.replace('क्ष', 'I')
    text = text.replace('त्र', '«')
    text = text.replace('ज्ञ', 'Ü')

    mapping = {
        'क': 's',
        'ख': 'v',
        'ग': 'u',
        'घ': '3',
        'ङ': 'ª',
        'च': 'r',
        'छ': '5',
        'ज': 'h',
        'ट': '6',
        'ठ': '7',
        'ड': '8',
        'ढ': '9',
        'त': 't',
        'थ': 'y',
        'द': 'b',
        'ध': 'w',
        'न': 'g',
        'प': 'k',
        'फ': 'f',
        'ब': 'a',
        'भ': 'e',
        'म': 'd',
        'य': 'o',
        'र': '/',
        'ल': 'n',
        'व': 'j',
        'स': ';',
        'ह': 'x',
        'ा': 'f',
        'ि': 'l',
        'ी': 'L',
        'ु': ']',
        'ू': '}',
        'ै': 'G',
        'ं': '+',
        'ः': 'M',
        '्': '\\',
    }

    result = ''
    for ch in text:
        result += mapping.get(ch, ch)

    return result


def process_paragraphs(paragraphs):
    for para in paragraphs:
        for run in para.runs:
            run.text = unicode_to_preeti(run.text)
            run.font.name = "Preeti"


def process_tables(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)


def process_shapes(doc):
    try:
        for shape in doc.inline_shapes:
            if hasattr(shape, "text_frame"):
                process_paragraphs(shape.text_frame.paragraphs)
    except Exception:
        pass


doc = Document(file_path)

process_paragraphs(doc.paragraphs)
process_tables(doc.tables)
process_shapes(doc)   # ✅ now used

doc.save("output_preeti.docx")